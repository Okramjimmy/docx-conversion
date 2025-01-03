import PyPDF2
from docx import Document
import os

def pdf_to_docx(pdf_path, docx_path):
    # Check if the PDF file exists
    if not os.path.exists(pdf_path):
        print(f"Error: The file {pdf_path} does not exist.")
        return

    try:
        # Initialize a PDF reader object
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Create a new DOCX document
            doc = Document()

            # Iterate through each page of the PDF
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()

                # If there's any text on the page, add it to the DOCX document
                if text:
                    doc.add_paragraph(text)
                    doc.add_paragraph("\n")  # Add a line break between pages
                else:
                    print(f"Warning: No text found on page {page_num + 1}")

            # Save the DOCX file
            doc.save(docx_path)
        
        print(f"PDF converted to DOCX successfully: {docx_path}")

    except Exception as e:
        print(f"Error during conversion: {e}")


# Example Usage
pdf_path = "GHC_Tools_new.pdf"  # Your input PDF file path
docx_path = "GHC_Tools_new.docx"  # Desired output DOCX file path

pdf_to_docx(pdf_path, docx_path)
